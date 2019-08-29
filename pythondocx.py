#这一篇学习python write docx文档
#当你了解了基本的程序底层逻辑之后，其他的都不过是熟悉api，读文档的事情了，所以底层的写代码逻辑要熟练
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH #段落的对齐方式
document = docx.Document()  #python对大小写敏感
print(dir(document))
def writedocx():
    document.add_heading('这是一个标题',level = 1)
    p1 = document.add_paragraph('这是一个段落，这是一个段落，这是一个段落，这是一个段落')

    p2 = p1.insert_paragraph_before('在一个段落的前面插入一段代码')

    document.add_page_break()  #另起一页，添加分页符
    
    p2 = document.add_paragraph('这是一个无序列表1',style = 'List Bullet')
    p2 = document.add_paragraph('这是一个无序列表2',style = 'List Bullet')
    p2 = document.add_paragraph('这是一个无序列表3',style = 'List Bullet')
    p2 = document.add_paragraph('这是一个无序列表4',style = 'List Bullet')
    p2 = document.add_paragraph('这是一个有序列表1',style = 'List Number')
    p2 = document.add_paragraph('这是一个有序列表2',style = 'List Number')
    p2 = document.add_paragraph('这是一个有序列表3',style = 'List Number')
    p2 = document.add_paragraph('这是一个有序列表4',style = 'List Number')
    document.add_heading('这是第二页的标题',level = 2)
    p3 = document.add_paragraph('改变段落的样式，改变段落的样式，改变段落的样式')
    p3.add_run('可以使用add_run来给段落添加样式，加粗，改变,加粗样式').bold = True
    run = p3.add_run('斜体')
    run.italic = True
    run.style = 'Emphasis'
    table = document.add_table(1,3)  #创建一个表格
    #cell = table.cell(0,0)   #访问某一个单元格 cell是细胞的意思，一个细胞就是一个基本的单元格子
    #cell.text = 'name,age,weight'
    row_cells = table.rows[0].cells
    row_cells[0].text = '姓名'
    row_cells[1].text = '性别'
    row_cells[2].text = '年龄'
    items = (
        ('文孝礼','男','29'),
        ('小明','男','29'),
        ('小红','女','29')
    )
    for item in items:
        cells = table.add_row().cells  #添加行
        cells[0].text = item[0]
        cells[1].text = item[1]
        cells[2].text = item[2]
    table.style = 'LightShading-Accent1'  #指定表格的样式，可以从doc文件里面选择你需要的样式
    document.add_picture('./img/1.jpg',width=Inches(6.0))  #添加图片
    document.add_page_break()
    document.add_heading('这一节主要讲解如何给段落添加样式',level = 2)
    para = document.add_paragraph('段落可以有各种各样的对齐方式')
    paragraph_format = para.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2 = document.add_paragraph('段落缩进')
    paragraph_format = para2.paragraph_format
    paragraph_format.left_indent = Inches(0.5)
    document.save('./docdir/demo.docx')
    #访问就研究到这里了，剩下的待有需求的时候继续研究
    print('成功生成了一个docx文件')

def main():
    writedocx()
if __name__ == '__main__':
    main()
    







